# ====================================================================
# JARVIS OMEGA — Documents Plugin
# ====================================================================
"""
Phase 8 seed plugin: PDF / DOCX / XLSX read + write.

Optional dependencies: ``pypdf``, ``python-docx``, ``openpyxl``. Each tool
imports lazily and raises a helpful error if its library is missing.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

from backend.tools import tool, RiskTier


# --------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------

@tool(
    name="docs.read_pdf",
    description="Extract text from a PDF file. Returns the concatenated text of every page.",
    parameters={
        "type": "object",
        "properties": {"path": {"type": "string"}},
        "required": ["path"],
    },
    risk_tier=RiskTier.TIER_0_OBSERVE,
    category="documents",
)
async def docs_read_pdf(path: str) -> Dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("pypdf is not installed — add `pypdf` to requirements.txt") from e
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return {"pages": len(pages), "text": "\n\n".join(pages)}


# --------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------

@tool(
    name="docs.read_docx",
    description="Extract paragraph text from a .docx file.",
    parameters={
        "type": "object",
        "properties": {"path": {"type": "string"}},
        "required": ["path"],
    },
    risk_tier=RiskTier.TIER_0_OBSERVE,
    category="documents",
)
async def docs_read_docx(path: str) -> Dict[str, Any]:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return {"paragraphs": len(paragraphs), "text": "\n".join(paragraphs)}


@tool(
    name="docs.write_docx",
    description="Create a new .docx file with the given paragraphs.",
    parameters={
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "paragraphs": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["path", "paragraphs"],
    },
    risk_tier=RiskTier.TIER_1_REVERSIBLE,
    category="documents",
)
async def docs_write_docx(path: str, paragraphs: List[str]) -> Dict[str, Any]:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)
    return {"path": str(Path(path).resolve()), "paragraphs": len(paragraphs)}


# --------------------------------------------------------------------
# XLSX
# --------------------------------------------------------------------

@tool(
    name="docs.read_xlsx",
    description="Read an .xlsx workbook. Returns a dict mapping sheet name -> list of row lists.",
    parameters={
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "max_rows": {"type": "integer", "default": 1000},
        },
        "required": ["path"],
    },
    risk_tier=RiskTier.TIER_0_OBSERVE,
    category="documents",
)
async def docs_read_xlsx(path: str, max_rows: int = 1000) -> Dict[str, Any]:
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise RuntimeError("openpyxl not installed") from e
    wb = load_workbook(path, read_only=True, data_only=True)
    sheets: Dict[str, List[List[Any]]] = {}
    for ws in wb.worksheets:
        rows: List[List[Any]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            rows.append(list(row))
        sheets[ws.title] = rows
    wb.close()
    return {"sheets": sheets}


@tool(
    name="docs.write_xlsx",
    description="Create a new .xlsx file with the given sheets.",
    parameters={
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "sheets": {
                "type": "object",
                "description": "Dict mapping sheet name -> 2D array of cell values.",
                "additionalProperties": {"type": "array"},
            },
        },
        "required": ["path", "sheets"],
    },
    risk_tier=RiskTier.TIER_1_REVERSIBLE,
    category="documents",
)
async def docs_write_xlsx(path: str, sheets: Dict[str, List[List[Any]]]) -> Dict[str, Any]:
    try:
        from openpyxl import Workbook
    except ImportError as e:
        raise RuntimeError("openpyxl not installed") from e
    wb = Workbook()
    # Remove the default sheet; we'll create our own.
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name[:31])  # Excel sheet name cap
        for row in rows:
            ws.append(row)
    wb.save(path)
    return {"path": str(Path(path).resolve()), "sheets": list(sheets.keys())}


PLUGIN_NAME = "documents"
PLUGIN_VERSION = "1.0.0"
PLUGIN_DESCRIPTION = "PDF / DOCX / XLSX read + write."
